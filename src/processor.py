import PyPDF2
import docx
import io


def process_file(uploaded_file) -> str:
    """Process uploaded file and extract ALL text."""
    file_name = uploaded_file.name.lower()

    if file_name.endswith(".pdf"):
        return extract_pdf_text(uploaded_file)
    elif file_name.endswith(".txt"):
        return extract_txt_text(uploaded_file)
    elif file_name.endswith(".docx"):
        return extract_docx_text(uploaded_file)
    else:
        raise ValueError(f"Unsupported file format: {file_name}")


def extract_pdf_text(uploaded_file) -> str:
    """Extract ALL text from PDF — all pages."""
    try:
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        total_pages = len(pdf_reader.pages)
        text_parts = []

        for page_num in range(total_pages):
            try:
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
            except Exception as e:
                text_parts.append(f"--- Page {page_num + 1} --- [Could not extract: {e}]")

        full_text = "\n\n".join(text_parts)

        if not full_text.strip():
            raise ValueError("No text could be extracted from PDF. File may be scanned or image-based.")

        return full_text

    except Exception as e:
        raise ValueError(f"PDF processing error: {e}")


def extract_txt_text(uploaded_file) -> str:
    """Extract text from TXT file."""
    try:
        content = uploaded_file.read()
        for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
            try:
                return content.decode(encoding)
            except:
                continue
        return content.decode("utf-8", errors="ignore")
    except Exception as e:
        raise ValueError(f"TXT processing error: {e}")


def extract_docx_text(uploaded_file) -> str:
    """Extract ALL text from DOCX including tables."""
    try:
        doc = docx.Document(uploaded_file)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts)
        if not full_text.strip():
            raise ValueError("No text found in DOCX file.")
        return full_text

    except Exception as e:
        raise ValueError(f"DOCX processing error: {e}")


def chunk_text(text: str, chunk_size: int = 3000) -> list:
    """Split text into chunks for AI processing."""
    words = text.split()
    chunks = []
    current_chunk = []
    current_size = 0

    for word in words:
        current_chunk.append(word)
        current_size += len(word) + 1
        if current_size >= chunk_size:
            chunks.append(" ".join(current_chunk))
            current_chunk = []
            current_size = 0

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks